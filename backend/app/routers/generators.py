import io
import zipfile
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from docx import Document as DocxDocument

from backend.app.db.database import get_db
from backend.app.db.models import User, Skill, Project, Certificate, Internship, Achievement, Timeline, Notification
from backend.app.routers.auth import get_current_user

router = APIRouter(prefix="/api/generators", tags=["Generators"])

@router.get("/resume/docx")
def generate_resume_docx(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    # 1. Fetch user data
    skills = db.query(Skill).filter(Skill.user_id == current_user.id).all()
    projects = db.query(Project).filter(Project.user_id == current_user.id).all()
    certs = db.query(Certificate).filter(Certificate.user_id == current_user.id).all()
    internships = db.query(Internship).filter(Internship.user_id == current_user.id).all()
    achievements = db.query(Achievement).filter(Achievement.user_id == current_user.id).all()

    # 2. Build DOCX
    doc = DocxDocument()
    
    # Title / Header
    name = current_user.full_name or "Student Profile"
    title_p = doc.add_paragraph()
    run = title_p.add_run(name)
    run.font.size = 24 * 12700  # Pt size (roughly)
    run.bold = True
    
    contact_p = doc.add_paragraph(f"Email: {current_user.email} | Powered by MemoryVerse AI")
    contact_p.alignment = 1 # Center
    
    doc.add_heading("Education & Internships", level=1)
    for intern in internships:
        doc.add_heading(f"{intern.role} – {intern.organization}", level=2)
        doc.add_paragraph(f"Duration: {intern.start_date} - {intern.end_date or 'Present'}")
        doc.add_paragraph(intern.description or "")
        if intern.responsibilities:
            doc.add_paragraph(f"Key Responsibilities:\n{intern.responsibilities}")

    doc.add_heading("Skills", level=1)
    skills_list = [s.name for s in skills]
    doc.add_paragraph(", ".join(skills_list) if skills_list else "No skills extracted yet.")

    doc.add_heading("Projects", level=1)
    for proj in projects:
        doc.add_heading(proj.name, level=2)
        doc.add_paragraph(f"Technologies: {proj.technologies}")
        doc.add_paragraph(proj.description or "")
        if proj.url:
            doc.add_paragraph(f"Project URL: {proj.url}")

    doc.add_heading("Certifications", level=1)
    for cert in certs:
        doc.add_paragraph(f"• {cert.name} (Issued by: {cert.authority}, Date: {cert.date})")

    doc.add_heading("Achievements", level=1)
    for ach in achievements:
        doc.add_paragraph(f"• {ach.title} (Organized by: {ach.organization or 'N/A'}, Date: {ach.date})")

    # Save to memory stream
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    # Notification
    notif = Notification(
        user_id=current_user.id,
        type="Resume Generated",
        message="Your ATS-friendly DOCX Resume was successfully generated."
    )
    db.add(notif)
    db.commit()

    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={name.replace(' ', '_')}_Resume.docx"}
    )

@router.get("/portfolio/export")
def export_portfolio_website(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    skills = db.query(Skill).filter(Skill.user_id == current_user.id).all()
    projects = db.query(Project).filter(Project.user_id == current_user.id).all()
    certs = db.query(Certificate).filter(Certificate.user_id == current_user.id).all()
    timeline = db.query(Timeline).filter(Timeline.user_id == current_user.id).order_by(Timeline.year.desc()).all()
    
    skills_json = [{"name": s.name, "category": s.category, "proficiency": s.proficiency} for s in skills]
    projects_json = [{"name": p.name, "description": p.description, "technologies": p.technologies, "url": p.url} for p in projects]
    certs_json = [{"name": c.name, "authority": c.authority, "date": c.date} for c in certs]
    timeline_json = [{"year": t.year, "event_title": t.event_title, "event_type": t.event_type, "description": t.description} for t in timeline]
    
    # Build Portfolio Static Page
    name = current_user.full_name or "Developer Profile"
    email = current_user.email
    
    portfolio_html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{name} - Portfolio Portfolio</title>
    <script src="https://cdn.tailwindcss.com"></script>
    <link href="https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;600;800&display=swap" rel="stylesheet">
    <style>
        body {{
            font-family: 'Outfit', sans-serif;
            background: linear-gradient(135deg, #0f172a 0%, #1e1b4b 100%);
        }}
    </style>
</head>
<body class="text-slate-100 min-h-screen">
    <header class="max-w-6xl mx-auto px-6 py-12 flex justify-between items-center">
        <h1 class="text-3xl font-extrabold bg-gradient-to-r from-blue-400 to-purple-500 bg-clip-text text-transparent">MemoryVerse</h1>
        <a href="mailto:{email}" class="px-6 py-2.5 rounded-full bg-indigo-600 hover:bg-indigo-700 transition font-medium">Contact Me</a>
    </header>

    <main class="max-w-6xl mx-auto px-6 space-y-24 pb-24">
        <!-- Hero Section -->
        <section class="text-center py-16 space-y-6">
            <h2 class="text-5xl md:text-7xl font-extrabold tracking-tight">Hi, I'm <span class="bg-gradient-to-r from-cyan-400 to-indigo-500 bg-clip-text text-transparent">{name}</span></h2>
            <p class="text-xl text-slate-400 max-w-2xl mx-auto">Welcome to my AI-curated student portfolio. Here you can explore my certificates, recent projects, skills, and academic achievements.</p>
        </section>

        <!-- Skills Section -->
        <section class="space-y-8">
            <h3 class="text-3xl font-bold border-b border-indigo-900 pb-4">My Skills</h3>
            <div class="flex flex-wrap gap-3" id="skills-container"></div>
        </section>

        <!-- Projects Section -->
        <section class="space-y-8">
            <h3 class="text-3xl font-bold border-b border-indigo-900 pb-4">Featured Projects</h3>
            <div class="grid md:grid-cols-2 gap-6" id="projects-container"></div>
        </section>

        <!-- Digital Timeline Section -->
        <section class="space-y-8">
            <h3 class="text-3xl font-bold border-b border-indigo-900 pb-4">Academic & Professional Journey</h3>
            <div class="space-y-6 border-l-2 border-indigo-600 pl-6 ml-4" id="timeline-container"></div>
        </section>
    </main>

    <footer class="text-center py-8 text-slate-500 border-t border-indigo-950">
        <p>&copy; 2026 {name}. Generated using MemoryVerse AI.</p>
    </footer>

    <script>
        const data = {{
            skills: {skills_json},
            projects: {projects_json},
            certs: {certs_json},
            timeline: {timeline_json}
        }};

        // Render Skills
        const skillsContainer = document.getElementById("skills-container");
        if (data.skills.length === 0) {{
            skillsContainer.innerHTML = "<p class='text-slate-400'>No skills added yet.</p>";
        }} else {{
            data.skills.forEach(skill => {{
                const el = document.createElement("span");
                el.className = "px-4 py-2 bg-indigo-950 border border-indigo-800 rounded-full text-indigo-300 font-semibold text-sm";
                el.innerText = skill.name;
                skillsContainer.appendChild(el);
            }});
        }}

        // Render Projects
        const projectsContainer = document.getElementById("projects-container");
        if (data.projects.length === 0) {{
            projectsContainer.innerHTML = "<p class='text-slate-400'>No projects added yet.</p>";
        }} else {{
            data.projects.forEach(p => {{
                const card = document.createElement("div");
                card.className = "p-6 rounded-2xl bg-indigo-950/40 border border-indigo-900/60 space-y-3";
                card.innerHTML = `
                    <h4 class="text-xl font-bold text-slate-100">${{p.name}}</h4>
                    <p class="text-slate-400 text-sm">${{p.description || ''}}</p>
                    <div class="flex justify-between items-center pt-2">
                        <span class="text-xs text-indigo-400 font-mono">${{p.technologies || ''}}</span>
                        ${{p.url ? `<a href="${{p.url}}" target="_blank" class="text-indigo-300 hover:underline text-sm font-medium">View Project &rarr;</a>` : ''}}
                    </div>
                `;
                projectsContainer.appendChild(card);
            }});
        }}

        // Render Timeline
        const timelineContainer = document.getElementById("timeline-container");
        if (data.timeline.length === 0) {{
            timelineContainer.innerHTML = "<p class='text-slate-400'>Timeline is empty.</p>";
        }} else {{
            data.timeline.forEach(item => {{
                const wrapper = document.createElement("div");
                wrapper.className = "relative space-y-2";
                wrapper.innerHTML = `
                    <div class="absolute -left-[31px] top-1.5 w-4 h-4 rounded-full bg-indigo-500 ring-4 ring-indigo-950"></div>
                    <span class="text-indigo-400 font-bold text-sm">${{item.year}}</span>
                    <h4 class="text-lg font-bold text-slate-100">${{item.event_title}}</h4>
                    <p class="text-slate-400 text-sm">${{item.description || ''}}</p>
                `;
                timelineContainer.appendChild(wrapper);
            }});
        }}
    </script>
</body>
</html>"""

    # Zip files together
    zip_stream = io.BytesIO()
    with zipfile.ZipFile(zip_stream, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        zip_file.writestr("index.html", portfolio_html)
        # Adding a small readme for the user
        zip_file.writestr("README.txt", f"MemoryVerse AI Portfolio Export\n===============================\nDouble-click index.html to launch your portfolio page locally!")
        
    zip_stream.seek(0)
    
    return StreamingResponse(
        zip_stream,
        media_type="application/zip",
        headers={"Content-Disposition": f"attachment; filename={name.replace(' ', '_')}_Portfolio.zip"}
    )
